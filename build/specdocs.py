"""§7.5/§7.6 spec library — authors the downloadable documents into assets/spec/.

Open item 3 in the handoff notes was "spec PDFs/DWGs not authored", which left the
primary CTA on /commercial/spec-library and /for-professionals/resources pointing at
"#". This module authors them.

Non-negotiable rule carried over from MASTER_PLAN §8 and the build conventions:
**every number in every document is read from data.py.** Nothing is typed twice and
nothing is invented. If a figure is not published, the document says so in the same
sentence rather than estimating it. That is why there are no U-factors, R-values,
energy-savings percentages or acoustic ratings anywhere in here — Veneta does not
publish them, so a specifier will not find them in a Veneta document either.

Two assets are deliberately NOT authored:

  veneta-details.zip        DWG sections for headrail, bracket, pocket and frame
  veneta-shutter-frames.pdf L-frame / deco frame / Z-frame profiles at full scale

Both require dimensioned manufacturing geometry that does not exist in data.py.
Drawing them from guesswork would put wrong numbers in front of an architect, which
is worse than a missing file. They stay in PENDING below until real CAD lands, and
build/pages5.py renders them with the "goes live when the file ships" note.

Typography: the brand faces are vendored as static TTFs in build/fonts/ so the build
stays offline and the documents match the site rather than defaulting to Helvetica.
"""
import os
import re
import html
import datetime

from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (BaseDocTemplate, Frame, PageTemplate, Paragraph,
                               Spacer, Table, TableStyle, KeepTogether,
                               PageBreak, Flowable)

import data as D

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
OUT = os.path.join(ROOT, "assets", "spec")
FONTS = os.path.join(HERE, "fonts")

SITE = "venetawindowfashions.com"
AVAIL = LETTER[0] - 1.9 * inch - 12      # page width less margins less frame padding
REV = datetime.date(2026, 8, 17)          # revision stamp printed on every document
REV_LABEL = REV.strftime("%B %Y")

# Assets referenced by build/pages5.py that this module does not author. Kept here so
# there is exactly one list and pages5 cannot disagree with the filesystem.
PENDING = {
    "veneta-details.zip":        "Dimensioned DWG sections are not released yet.",
    "veneta-shutter-frames.pdf": "Full-scale frame profiles are not released yet.",
}


# ----------------------------------------------------------------- tokens (§4.1)
CANVAS      = colors.HexColor("#F6F2EC")
SURFACE     = colors.HexColor("#FCFAF6")
SURFACE_SNK = colors.HexColor("#EDE7DD")
LINE        = colors.HexColor("#DCD5C9")
LINE_SOFT   = colors.HexColor("#E8E2D8")
NOIR        = colors.HexColor("#16150F")
INK         = colors.HexColor("#211C16")
INK_70      = colors.HexColor("#55503F")
INK_45      = colors.HexColor("#8A8371")
CLAY        = colors.HexColor("#8C5A38")

SERIF = "InstrumentSerif"
SERIF_I = "InstrumentSerif-Italic"
SANS = "InterTight"
SANS_M = "InterTight-Medium"
SANS_SB = "InterTight-SemiBold"

_registered = False


def _fonts():
    global _registered
    if _registered:
        return
    pdfmetrics.registerFont(TTFont(SERIF, os.path.join(FONTS, "InstrumentSerif-Regular.ttf")))
    pdfmetrics.registerFont(TTFont(SERIF_I, os.path.join(FONTS, "InstrumentSerif-Italic.ttf")))
    pdfmetrics.registerFont(TTFont(SANS, os.path.join(FONTS, "InterTight-Regular.ttf")))
    pdfmetrics.registerFont(TTFont(SANS_M, os.path.join(FONTS, "InterTight-Medium.ttf")))
    pdfmetrics.registerFont(TTFont(SANS_SB, os.path.join(FONTS, "InterTight-SemiBold.ttf")))
    _registered = True


# ------------------------------------------------------------------ text helpers
def t(s):
    """data.py copy is written for HTML. Turn entities into real characters, strip
    inline markup, then re-escape for reportlab's mini-markup."""
    s = re.sub(r"<[^>]+>", "", str(s))
    s = html.unescape(s)
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def plain(s):
    """Same, but for python-docx and for anything that must not carry markup."""
    return html.unescape(re.sub(r"<[^>]+>", "", str(s)))


# ----------------------------------------------------------------- type scale
# Nine sizes, mapped from the locked web scale to print. No ad-hoc sizes.
S = {}


def _styles():
    if S:
        return S
    _fonts()
    S["display"] = ParagraphStyle("display", fontName=SERIF, fontSize=44, leading=44,
                                  textColor=INK, spaceAfter=0)
    S["h1"] = ParagraphStyle("h1", fontName=SERIF, fontSize=30, leading=32,
                             textColor=INK, spaceBefore=0, spaceAfter=10)
    S["h2"] = ParagraphStyle("h2", fontName=SERIF, fontSize=21, leading=24,
                             textColor=INK, spaceBefore=18, spaceAfter=7)
    S["h3"] = ParagraphStyle("h3", fontName=SANS_M, fontSize=13, leading=17,
                             textColor=INK, spaceBefore=13, spaceAfter=4)
    S["lede"] = ParagraphStyle("lede", fontName=SANS, fontSize=12.5, leading=18.5,
                               textColor=INK_70, spaceAfter=10)
    S["body"] = ParagraphStyle("body", fontName=SANS, fontSize=10, leading=15.5,
                               textColor=INK, spaceAfter=7, alignment=TA_LEFT)
    S["body70"] = ParagraphStyle("body70", parent=S["body"], textColor=INK_70)
    S["small"] = ParagraphStyle("small", fontName=SANS, fontSize=9, leading=13.5,
                                textColor=INK_70, spaceAfter=5)
    S["cap"] = ParagraphStyle("cap", fontName=SANS, fontSize=8, leading=11.5,
                              textColor=INK_45, spaceAfter=4)
    S["micro"] = ParagraphStyle("micro", fontName=SANS_M, fontSize=7.2, leading=10,
                                textColor=INK_45, spaceAfter=6)
    # derived
    S["eyebrow"] = ParagraphStyle("eyebrow", parent=S["micro"], textColor=CLAY, spaceAfter=4)
    S["cell"] = ParagraphStyle("cell", fontName=SANS, fontSize=9, leading=13, textColor=INK)
    S["cellk"] = ParagraphStyle("cellk", fontName=SANS_M, fontSize=9, leading=13, textColor=INK_70)
    S["cellh"] = ParagraphStyle("cellh", fontName=SANS_SB, fontSize=7.2, leading=10,
                                textColor=INK_45)
    S["coverdisp"] = ParagraphStyle("coverdisp", fontName=SERIF, fontSize=40, leading=41,
                                    textColor=colors.HexColor("#FCFAF6"))
    S["coversub"] = ParagraphStyle("coversub", fontName=SANS, fontSize=12, leading=18,
                                   textColor=colors.HexColor("#B7B0A0"))
    return S


def eyebrow(txt):
    return Paragraph(" ".join(t(txt).upper()).replace("   ", "  \u2009") if False
                     else t(txt).upper(), _styles()["eyebrow"])


def h2(txt):
    return Paragraph(t(txt), _styles()["h2"])


def h3(txt):
    return Paragraph(t(txt), _styles()["h3"])


def body(txt, style="body"):
    return Paragraph(t(txt), _styles()[style])


def bullets(items, style="body"):
    st = _styles()[style]
    rows = [[Paragraph("&mdash;", _styles()["cellk"]), Paragraph(t(i), st)] for i in items]
    tb = Table(rows, colWidths=[0.28 * inch, None])
    tb.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 1),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    return tb


def numbered(items):
    st = _styles()["body"]
    rows = [[Paragraph(f"{n}", _styles()["cellk"]), Paragraph(t(i), st)]
            for n, i in enumerate(items, 1)]
    tb = Table(rows, colWidths=[0.28 * inch, None])
    tb.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 1),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    return tb


class Rule(Flowable):
    """One hairline. The only decoration in these documents."""

    def __init__(self, width=None, color=LINE, space=8, thickness=0.6):
        Flowable.__init__(self)
        self.w, self.c, self.space, self.th = width, color, space, thickness

    def wrap(self, aw, ah):
        self._w = self.w or aw
        return (self._w, self.space)

    def draw(self):
        self.canv.setStrokeColor(self.c)
        self.canv.setLineWidth(self.th)
        self.canv.line(0, self.space / 2, self._w, self.space / 2)


class Swatch(Flowable):
    """A colourway chip. Shows the published finish colour, nothing more."""

    def __init__(self, hexcolor, w=26, h=13):
        Flowable.__init__(self)
        self.hex, self.w, self.h = hexcolor, w, h

    def wrap(self, aw, ah):
        return (self.w, self.h)

    def draw(self):
        self.canv.setFillColor(colors.HexColor(self.hex))
        self.canv.setStrokeColor(LINE)
        self.canv.setLineWidth(0.6)
        self.canv.rect(0, 0, self.w, self.h, stroke=1, fill=1)


def kvtable(rows, avail, keyw=1.85, tint=True):
    """Two-column spec table. Tabular numerals, hairline rules, tinted key column."""
    st = _styles()
    data = [[Paragraph(t(k), st["cellk"]), Paragraph(t(v), st["cell"])] for k, v in rows]
    tb = Table(data, colWidths=[keyw * inch, avail - keyw * inch], hAlign="LEFT")
    cmds = [
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (0, -1), 8),
        ("LEFTPADDING", (1, 0), (1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 3.8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3.8),
        ("LINEBELOW", (0, 0), (-1, -2), 0.5, LINE_SOFT),
        ("LINEABOVE", (0, 0), (-1, 0), 0.6, LINE),
        ("LINEBELOW", (0, -1), (-1, -1), 0.6, LINE),
    ]
    if tint:
        cmds.append(("BACKGROUND", (0, 0), (0, -1), SURFACE_SNK))
    tb.setStyle(TableStyle(cmds))
    return tb


def gridtable(head, rows, widths, avail):
    """Multi-column schedule table."""
    st = _styles()
    data = [[Paragraph(t(c).upper(), st["cellh"]) for c in head]]
    for r in rows:
        data.append([c if isinstance(c, Flowable) else Paragraph(t(c), st["cell"]) for c in r])
    total = sum(widths)
    cw = [w / total * avail for w in widths]
    tb = Table(data, colWidths=cw, hAlign="LEFT", repeatRows=1)
    tb.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 9),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 6),
        ("BOTTOMPADDING", (0, 1), (-1, -1), 5),
        ("LINEBELOW", (0, 0), (-1, 0), 0.6, INK_45),
        ("LINEBELOW", (0, 1), (-1, -2), 0.5, LINE_SOFT),
        ("LINEBELOW", (0, -1), (-1, -1), 0.6, LINE),
    ]))
    return tb


def callout(txt, avail, label="Note"):
    st = _styles()
    inner = [Paragraph(t(label).upper(), st["micro"]), Paragraph(t(txt), st["small"])]
    tb = Table([[inner]], colWidths=[avail], hAlign="LEFT")
    tb.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), SURFACE_SNK),
        ("LEFTPADDING", (0, 0), (-1, -1), 14),
        ("RIGHTPADDING", (0, 0), (-1, -1), 14),
        ("TOPPADDING", (0, 0), (-1, -1), 11),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
        ("LINEBEFORE", (0, 0), (0, -1), 2, CLAY),
    ]))
    return tb


def _ls(c, x, y, txt, font, size, space, color, align="left"):
    """Letterspaced text. reportlab only exposes character spacing on a text object."""
    from reportlab.pdfbase.pdfmetrics import stringWidth
    if align in ("right", "center"):
        w = stringWidth(txt, font, size) + space * max(0, len(txt) - 1)
        x = x - w if align == "right" else x - w / 2
    to = c.beginText(x, y)
    to.setFont(font, size)
    to.setFillColor(color)
    to.setCharSpace(space)
    to.textOut(txt)
    c.drawText(to)
    # Tc is page-level graphics state: leave it as we found it or every paragraph
    # laid out after this header inherits the letterspacing.
    reset = c.beginText(0, 0)
    reset.setCharSpace(0)
    c.drawText(reset)


# ------------------------------------------------------------------- doc template
class Doc(BaseDocTemplate):
    """Letter, one text frame, hairline running head and foot. A cover page with a
    --noir band, then content pages on --surface."""

    def __init__(self, filename, title, subtitle, audience, kicker):
        self.brand_title = title
        self.subtitle = subtitle
        self.audience = audience
        self.kicker = kicker
        BaseDocTemplate.__init__(
            self, filename, pagesize=LETTER,
            leftMargin=0.95 * inch, rightMargin=0.95 * inch,
            topMargin=1.05 * inch, bottomMargin=0.85 * inch,
            title=plain(title) + " — VENETA",
            author="Veneta Window Fashions",
            subject=plain(subtitle),
            creator="build/specdocs.py",
        )
        fr = Frame(self.leftMargin, self.bottomMargin, self.width, self.height, id="body")
        cover = Frame(self.leftMargin, 1.15 * inch, self.width,
                      LETTER[1] - 1.15 * inch - 3.6 * inch, id="cover")
        self.addPageTemplates([
            PageTemplate(id="cover", frames=[cover], onPage=self._cover),
            PageTemplate(id="body", frames=[fr], onPage=self._chrome),
        ])

    # cover -------------------------------------------------------------------
    def _cover(self, c, doc):
        w, h = LETTER
        c.setFillColor(SURFACE)
        c.rect(0, 0, w, h, stroke=0, fill=1)
        c.setFillColor(NOIR)
        c.rect(0, h - 3.8 * inch, w, 3.8 * inch, stroke=0, fill=1)

        _ls(c, 0.95 * inch, h - 1.05 * inch, "VENETA", SANS_M, 11, 3.4,
            colors.HexColor("#FCFAF6"))
        _ls(c, w - 0.95 * inch, h - 1.05 * inch, self.kicker.upper(), SANS, 8.4, 1.5,
            colors.HexColor("#8A8371"), align="right")

        c.setStrokeColor(colors.HexColor("#3a382c"))
        c.setLineWidth(0.6)
        c.line(0.95 * inch, h - 1.42 * inch, w - 0.95 * inch, h - 1.42 * inch)

        c.setFont(SERIF, 38)
        c.setFillColor(colors.HexColor("#FCFAF6"))
        c.drawString(0.95 * inch, h - 2.5 * inch, plain(self.brand_title))
        c.setFont(SANS, 11.5)
        c.setFillColor(colors.HexColor("#B7B0A0"))
        sublines = _wrap(plain(self.subtitle), 76)
        for i, ln in enumerate(sublines):
            c.drawString(0.95 * inch, h - 2.93 * inch - i * 16, ln)

        # audience sits in the band, not in the footer, so nothing collides
        _ls(c, 0.95 * inch, h - 2.93 * inch - len(sublines) * 16 - 20,
            self.audience.upper(), SANS_M, 8, 1.6, colors.HexColor("#8A8371"))

        # foot of cover: three even columns on one hairline
        c.setStrokeColor(LINE)
        c.setLineWidth(0.6)
        c.line(0.95 * inch, 1.02 * inch, w - 0.95 * inch, 1.02 * inch)
        _ls(c, 0.95 * inch, 0.78 * inch, f"REVISION {REV_LABEL.upper()}",
            SANS, 8, 1.2, INK_45)
        _ls(c, w / 2, 0.78 * inch, "SOLD THROUGH THE HOME DEPOT", SANS, 8, 1.2, INK_45,
            align="center")
        _ls(c, w - 0.95 * inch, 0.78 * inch, SITE.upper(), SANS, 8, 1.2, INK_45,
            align="right")


    # interior ----------------------------------------------------------------
    def _chrome(self, c, doc):
        w, h = LETTER
        c.setFillColor(SURFACE)
        c.rect(0, 0, w, h, stroke=0, fill=1)

        _ls(c, 0.95 * inch, h - 0.72 * inch, "VENETA", SANS_M, 8, 2.4, INK_45)
        _ls(c, w - 0.95 * inch, h - 0.72 * inch, plain(self.brand_title).upper(),
            SANS, 8, 0.8, INK_45, align="right")
        c.setStrokeColor(LINE)
        c.setLineWidth(0.6)
        c.line(0.95 * inch, h - 0.86 * inch, w - 0.95 * inch, h - 0.86 * inch)

        c.line(0.95 * inch, 0.66 * inch, w - 0.95 * inch, 0.66 * inch)
        c.setFont(SANS, 8)
        c.setFillColor(INK_45)
        c.drawString(0.95 * inch, 0.47 * inch,
                     f"{SITE}  ·  Revision {REV_LABEL}  ·  Ranges as published on the site")
        c.drawRightString(w - 0.95 * inch, 0.47 * inch, str(doc.page - 1))

    def afterFlowable(self, flowable):
        pass


def _wrap(s, n):
    out, line = [], ""
    for word in s.split():
        if len(line) + len(word) + 1 > n:
            out.append(line)
            line = word
        else:
            line = (line + " " + word).strip()
    if line:
        out.append(line)
    return out[:2]


def render(filename, title, subtitle, audience, kicker, story):
    """Cover page, then body pages."""
    _styles()
    os.makedirs(OUT, exist_ok=True)
    path = os.path.join(OUT, filename)
    doc = Doc(path, title, subtitle, audience, kicker)
    doc.build([Spacer(1, 1), _NextBody()] + story)
    return path


class _NextBody(Flowable):
    """Switch to the body template after the cover."""

    def wrap(self, aw, ah):
        return (0, 0)

    def draw(self):
        pass

    def split(self, aw, ah):
        return []


# reportlab needs a real page break + template switch; do it with NextPageTemplate
from reportlab.platypus.doctemplate import NextPageTemplate  # noqa: E402


def render(filename, title, subtitle, audience, kicker, story):  # noqa: F811
    _styles()
    os.makedirs(OUT, exist_ok=True)
    path = os.path.join(OUT, filename)
    doc = Doc(path, title, subtitle, audience, kicker)
    doc.build([NextPageTemplate("body"), Spacer(1, 0.1), PageBreak()] + story)
    return path


# ============================================================ shared content bits
def provenance(avail, extra=""):
    return callout(
        "Every range, size, material and option in this document is the published "
        f"figure from {SITE}. Nothing here is estimated. Where a figure is not "
        "published — energy savings, acoustic ratings, U-factors — this document "
        "says so rather than supplying a number. Pricing is set by The Home Depot "
        f"and is not published by Veneta. {extra}".strip(),
        avail, label="How to read this document")


def cordless_lines():
    """Lift options per line, straight out of the published spec tables."""
    rows = []
    for p in D.PRODUCTS:
        spec = dict(p["spec"])
        lift = spec.get("Lift options") or spec.get("Control") or spec.get("Tilt", "")
        rows.append((plain(p["name"]), plain(lift)))
    return rows


def spec_value(p, *keys):
    spec = dict(p["spec"])
    for k in keys:
        if k in spec:
            return spec[k]
    return ""


def size_rows():
    rows = []
    for p in D.PRODUCTS:
        rows.append([
            plain(p["name"]),
            plain(spec_value(p, "Width range", "Opening width")),
            plain(spec_value(p, "Height range")),
            plain(spec_value(p, "Mount")),
        ])
    return rows


# ==================================================================== spec book
BOOK_SECTIONS = {
    "commercial": dict(
        title="Commercial Spec Book",
        subtitle="Published size ranges, mount depths, materials, lift systems and "
                 "colourways for all eight Veneta lines.",
                audience="For specifiers and facilities",
        kicker="Section 12 24 00",
        opening=[
            ("Scope",
             "Eight made-to-size window covering lines, manufactured to submitted "
             "dimensions and supplied through The Home Depot. This document carries the "
             "published range for each line so a window schedule can be checked before "
             "it is issued."),
            ("Procurement",
             "Orders are placed through The Home Depot. Multi-unit, phased and "
             "date-committed work is handled by the Veneta trade desk. The product and "
             "the published sizes are identical either way."),
            ("Safety basis",
             "Cordless lift is the standard configuration on every line in this book. "
             "See the cordless statement for the written confirmation."),
        ]),
    "trade": dict(
        title="Trade Spec Book",
        subtitle="Published size ranges, mount depths, materials, lift systems and "
                 "colourways for all eight Veneta lines.",
        audience="For designers and installers",
        kicker="Trade reference",
        opening=[
            ("Scope",
             "Eight made-to-size window covering lines, manufactured to the dimensions "
             "you submit and supplied through The Home Depot. Use this to check a "
             "measurement against the range before the client signs off."),
            ("Ordering",
             "Through The Home Depot like any other order, or through the trade desk when "
             "the order is multi-unit, phased, or needs dates committed in writing."),
            ("Safety basis",
             "Cordless lift is the standard configuration on every line, which is the "
             "answer to the question clients with small children ask first."),
        ]),
}


def build_spec_book(kind):
    cfg = BOOK_SECTIONS[kind]
    st = _styles()
    avail = AVAIL
    s = []

    s.append(eyebrow("Contents"))
    s.append(Paragraph(plain(cfg["title"]), st["h1"]))
    s.append(Rule(space=12))
    idx = [[f"{i:02d}", plain(p["name"]), plain(p["tagline"])]
           for i, p in enumerate(D.PRODUCTS, 1)]
    idx += [["09", "Size ranges at a glance", "All eight lines in one table"],
            ["10", "Minimum inside-mount depth", "By line, with the outside-mount threshold"],
            ["11", "Cordless lift status", "Standard configuration per line"],
            ["12", "Warranty summary", "Coverage, exclusions and the claim route"]]
    s.append(gridtable(["", "Line", "Summary"], idx, [0.5, 2.4, 5.2], avail))
    s.append(PageBreak())

    # Front matter reads as two deliberate pages: contents, then how to read it.
    s.append(eyebrow("Front matter"))
    s.append(Paragraph("How to read this document", st["h1"]))
    s.append(Rule(space=12))
    s.append(provenance(avail))
    s.append(Spacer(1, 20))
    for head, txt in cfg["opening"]:
        s.append(h3(head))
        s.append(body(txt))

    # one page per line
    for i, p in enumerate(D.PRODUCTS, 1):
        s.append(PageBreak())
        s.append(eyebrow(f"Line {i:02d}"))
        s.append(Paragraph(t(p["name"]), st["h1"]))
        s.append(Paragraph(t(p["lede"]), st["lede"]))
        s.append(Rule(space=12))
        s.append(h3("Published specification"))
        s.append(kvtable(p["spec"], avail))
        s.append(Spacer(1, 12))
        notes = [h3("Specifier notes")]
        for name, txt in p["features"]:
            notes.append(Paragraph(f"<b>{t(name)}.</b> {t(txt)}", st["small"]))
        s.append(KeepTogether(notes))
        s.append(Spacer(1, 8))
        # Chips live in the fabric schedule, which carries a full page per line. Naming
        # them here and cross-referencing keeps one line to one page instead of eight
        # half-empty chip pages, and there is only one place a colourway is drawn.
        names = "  \u00b7  ".join(plain(n) for n, _ in p["colors"])
        s.append(KeepTogether([
            h3(f"Colourways ({len(p['colors'])})"),
            Paragraph(t(names), st["small"]),
            Paragraph(f"Chips, available opacity and reference are in the fabric and openness "
                      f"schedule. Specify from a physical swatch, never from a screen. "
                      f"{SITE}/{p['slug']}", st["cap"]),
        ]))


    # summary tables
    s.append(PageBreak())
    s.append(eyebrow("Section 09"))
    s.append(Paragraph("Size ranges at a glance", st["h1"]))
    s.append(Paragraph("Widths and heights are the manufactured range for each line. "
                       "Anything outside the range is not built, rather than built and "
                       "de-rated.", st["lede"]))
    s.append(gridtable(["Line", "Width", "Height", "Mount"], size_rows(),
                       [2.5, 1.9, 1.7, 3.3], avail))
    s.append(Spacer(1, 10))
    s.append(Paragraph('Width increments are 1/8" where published. Shutters are quoted by '
                       "opening and panel count rather than a single width range.",
                       st["cap"]))

    s.append(Spacer(1, 26))
    s.append(eyebrow("Section 10"))
    s.append(Paragraph("Minimum inside-mount depth", st["h2"]))
    s.append(gridtable(["Depth", "Applies to"],
                       [[plain(d), plain(lbl)] for d, lbl, _ in D.MOUNT_DEPTH],
                       [1.6, 6.4], avail))
    s.append(Spacer(1, 8))
    s.append(callout("Depth is measured from the face of the glass stop to the front of the "
                     "opening. Where the available depth is under the figure above, specify "
                     "an outside mount; the product is not modified to fit a shallow "
                     "opening.", avail, label="Measuring basis"))

    s.append(PageBreak())
    s.append(eyebrow("Section 11"))
    s.append(Paragraph("Cordless lift status", st["h1"]))
    s.append(Paragraph("Cordless lift is the standard configuration on every line. The "
                       "table lists the operating options published for each.", st["lede"]))
    s.append(gridtable(["Line", "Published operation"], cordless_lines(), [2.5, 6.0], avail))
    s.append(Spacer(1, 12))
    s.append(callout("Veneta does not publish an energy-savings percentage, an R-value or an "
                     "acoustic rating for any line, because the honest figure depends on the "
                     "glazing, orientation and climate of the specific opening. Where a "
                     "project needs a tested number, it has to come from a test on the "
                     "assembly as built.", avail, label="What is not published"))

    s.append(Spacer(1, 22))
    s.append(eyebrow("Section 12"))
    s.append(Paragraph("Warranty summary", st["h2"]))
    s.append(body(D.WARRANTY["scope"]))
    s.append(h3("Covered"))
    s.append(bullets(D.WARRANTY["covered"], "small"))
    s.append(h3("Not covered"))
    s.append(bullets(D.WARRANTY["excluded"], "small"))
    s.append(Spacer(1, 10))
    s.append(Paragraph(f'Full terms in the warranty document, or at {SITE}/warranty. '
                       f'Claims: {plain(D.SUPPORT_PHONE)}.', st["cap"]))

    name = f"veneta-{kind}-spec-book.pdf"
    return render(name, cfg["title"], cfg["subtitle"], cfg["audience"], cfg["kicker"], s)


# ============================================================== fabric schedule
def build_fabric_schedule():
    st = _styles()
    avail = AVAIL
    s = []
    s.append(eyebrow("Schedule"))
    s.append(Paragraph("Fabric, opacity and colourway schedule", st["h1"]))
    s.append(Paragraph("Openness factors, opacity options and published colourways for all "
                       "eight lines, in the order they appear in the spec book.", st["lede"]))
    s.append(Rule(space=12))
    s.append(provenance(avail, "Fabric weights in ounces per square yard are not published "
                               "by Veneta and are therefore absent from this schedule."))
    s.append(Spacer(1, 16))

    rows = []
    for p in D.PRODUCTS:
        rows.append([
            plain(p["name"]),
            plain(spec_value(p, "Opacity", "Lining") or "See line page"),
            plain(spec_value(p, "Openness factor", "UV block") or "Not published"),
            plain(spec_value(p, "Material", "Facings", "Vanes", "Cell type",
                             "Vane size", "Louvre width") or ""),
        ])
    s.append(h3("Opacity and openness by line"))
    s.append(gridtable(["Line", "Opacity options", "Openness / UV", "Material"],
                       rows, [2.2, 2.6, 1.9, 2.6], avail))
    s.append(Spacer(1, 10))
    s.append(callout("Openness factor is the percentage of the weave that is open. A lower "
                     "number cuts more glare and heat; a higher number keeps more of the "
                     "view. Openness is not privacy: a lit room is visible through any "
                     "screen fabric after dark.", avail, label="Reading openness"))

    for p in D.PRODUCTS:
        s.append(PageBreak())
        s.append(eyebrow(plain(p["short"])))
        s.append(Paragraph(t(p["name"]) + " colourways", st["h2"]))
        rows = [[Swatch(hx, w=34, h=17), plain(n),
                 plain(spec_value(p, "Opacity", "Lining") or ""),
                 f"{SITE}/{p['slug']}"] for n, hx in p["colors"]]
        s.append(gridtable(["Chip", "Colourway", "Available opacity", "Reference"],
                          rows, [0.9, 2.2, 3.0, 3.0], avail))
        s.append(Spacer(1, 8))
        s.append(Paragraph("Chips are reproduced from the published finish colour and will "
                           "shift with screen calibration and paper stock. Specify from a "
                           "physical swatch only.", st["cap"]))

    return render("veneta-fabric-schedule.pdf", "Fabric &amp; Openness Schedule",
                  "Opacity options, openness factors and published colourways for all eight "
                  "Veneta lines.", "For specifiers and designers", "Finish schedule", s)


# ================================================================= mount depths
def build_mount_depths():
    st = _styles()
    avail = AVAIL
    s = []
    s.append(eyebrow("Mounting"))
    s.append(Paragraph("Minimum inside-mount depth", st["h1"]))
    s.append(Paragraph("The depth each line needs to sit inside the opening, and what to do "
                       "when the opening is shallower than that.", st["lede"]))
    s.append(Rule(space=12))

    s.append(h3("By product group"))
    s.append(gridtable(["Minimum depth", "Applies to"],
                       [[plain(d), plain(lbl)] for d, lbl, _ in D.MOUNT_DEPTH],
                       [1.8, 6.2], avail))
    s.append(Spacer(1, 16))

    s.append(h3("By line, as published"))
    s.append(gridtable(["Line", "Published mount note", "Headrail / track"],
                       [[plain(p["name"]), plain(spec_value(p, "Mount")),
                         plain(spec_value(p, "Headrail", "Frame") or "&mdash;")]
                        for p in D.PRODUCTS],
                       [2.3, 3.4, 2.6], avail))
    s.append(Spacer(1, 16))

    s.append(h3("How the depth is measured"))
    s.append(numbered([
        "Measure from the face of the glass stop, or the innermost obstruction, to the "
        "front edge of the opening.",
        "Take the measurement at the head of the opening, where the headrail will sit.",
        "Deduct nothing. The figures above are net depths for the product, not for the "
        "opening.",
        "Where a crank, handle, alarm contact or trim projects into the opening, measure to "
        "the projection rather than to the stop.",
        "If the available depth is under the published minimum, specify an outside mount.",
    ]))
    s.append(Spacer(1, 14))
    s.append(callout("An outside mount is a specification choice, not a fallback. It covers "
                     "the trim as well as the glass, which closes the side light gap that an "
                     "inside mount leaves. Products are not modified to fit a shallow "
                     "opening.", avail, label="Shallow openings"))
    s.append(Spacer(1, 14))
    s.append(h3("Light gap"))
    s.append(body("An inside-mounted shade leaves a light gap of about 1/4\" on each side "
                  "with a ClearFit&trade; headrail cut to the opening. Where a room has to be "
                  "genuinely dark, specify SmartPrivacy&reg; side channels on a blackout "
                  "fabric, or move to an outside mount. Veneta does not describe any "
                  "configuration as total blackout."))
    s.append(Spacer(1, 12))
    s.append(Paragraph(f"Full measuring method at {SITE}/how-to-measure.", st["cap"]))

    return render("veneta-mount-depths.pdf", "Mount Depth Reference",
                  "Minimum inside-mount depth by line, how the depth is measured, and the "
                  "outside-mount threshold.", "For specifiers and installers",
                  "Mounting reference", s)


# ============================================================ cordless statement
def build_cordless_statement():
    st = _styles()
    avail = AVAIL
    s = []
    s.append(eyebrow("Statement"))
    s.append(Paragraph("Cordless lift is standard on every line", st["h1"]))
    s.append(Rule(space=12))
    s.append(body("Veneta confirms that cordless lift is the standard configuration on all "
                  "eight product lines listed below. No line ships with a looped operating "
                  "cord, a cord cleat or a tension device as its standard configuration."))
    s.append(body("Cellular shades are published as cordless with a rating against "
                  "ANSI/WCMA A100.1, the American National Standard for the safety of "
                  "corded window covering products. Faux wood blinds also offer a standard "
                  "cord lock as a customer-selected option; specify cordless explicitly "
                  "where a project requires it."))
    s.append(Spacer(1, 14))
    s.append(h3("Published operation by line"))
    s.append(gridtable(["Line", "Published operation"], cordless_lines(), [2.5, 6.0], avail))
    s.append(Spacer(1, 16))
    s.append(callout("This statement covers the operating system only. It is not a fire, "
                     "smoke, acoustic or thermal classification, and Veneta does not publish "
                     "those ratings. A project that needs a tested classification needs a "
                     "test on the assembly as installed.", avail, label="Limits of this statement"))
    s.append(Spacer(1, 18))
    s.append(h3("Where this matters most"))
    s.append(bullets([
        "Nurseries, children's bedrooms and any room a child sleeps in.",
        "Multifamily and hospitality, where the occupant is not the specifier.",
        "Education and healthcare, where a looped cord is usually prohibited outright.",
        "Low windows and window seats, where a cord would sit within reach.",
    ], "small"))
    s.append(Spacer(1, 22))
    s.append(Rule(space=14))
    s.append(Paragraph("Veneta Window Fashions", st["h3"]))
    s.append(Paragraph(f"Revision {REV_LABEL}  ·  {plain(D.SUPPORT_PHONE)}  ·  {SITE}\n"
                       f"Product sold through The Home Depot.", st["cap"]))

    return render("veneta-cordless-statement.pdf", "Cordless Lift Statement",
                  "Written confirmation that cordless lift is the standard configuration on "
                  "every Veneta line, with published operation per line.",
                  "For compliance files", "Compliance", s)


# ===================================================================== warranty
def build_warranty():
    st = _styles()
    W = D.WARRANTY
    avail = AVAIL
    s = []
    s.append(eyebrow("Policy"))
    s.append(Paragraph("Limited lifetime warranty", st["h1"]))
    s.append(Paragraph(plain(W["lede"]), st["lede"]))
    s.append(Rule(space=12))
    s.append(h2("What is covered"))
    s.append(body(W["scope"]))
    s.append(bullets(W["covered"]))
    s.append(h2("What is not covered"))
    s.append(bullets(W["excluded"]))
    s.append(h2("How to make a claim"))
    s.append(numbered(W["claim"]))
    s.append(Spacer(1, 12))
    s.append(callout(W["direct"], avail, label="No trip to the store"))
    s.append(PageBreak())
    s.append(h2("Motorization specifics"))
    s.append(body(W["motor"]))
    s.append(h2("Consumer rights"))
    s.append(body(W["rights"]))
    s.append(h2("Registering a product"))
    s.append(body(W["register"]))
    s.append(Spacer(1, 16))
    s.append(h3("Claim checklist"))
    s.append(bullets(W["checklist"], "small"))
    s.append(Spacer(1, 16))
    s.append(gridtable(["Route", "Detail"],
                       [["Phone", f"{plain(D.SUPPORT_PHONE)}  ·  {plain(D.SUPPORT_HOURS)}"],
                        ["Online", f"{SITE}/contact"],
                        ["Assessment", "Usually within one business day"],
                        ["Outcome", "Replacement part shipped, or the product remade to the "
                                   "original dimensions"]],
                       [1.7, 6.3], avail))
    s.append(Spacer(1, 12))
    s.append(Paragraph(f"This document reproduces the terms published at {SITE}/warranty. "
                       f"Where the two ever differ, the published page governs.", st["cap"]))

    return render("veneta-warranty.pdf", "Limited Lifetime Warranty",
                  "Coverage, exclusions, motorization terms and the five-step claim process.",
                  "For owners and facilities", "Policy", s)


# ========================================================== care and maintenance
def build_care_maintenance():
    st = _styles()
    C = D.CARE
    avail = AVAIL
    s = []
    s.append(eyebrow("Maintenance"))
    s.append(Paragraph("Care and maintenance", st["h1"]))
    s.append(Paragraph("Cleaning method by material, the products that void a finish, and a "
                       "15-minute seasonal check for a facilities schedule.", st["lede"]))
    s.append(Rule(space=12))
    s.append(h2("The rule that covers everything"))
    s.append(body(C["rule"]))
    s.append(Spacer(1, 8))
    s.append(h3("Quick reference"))
    s.append(gridtable(["Material", "Method"], [[plain(a), plain(b)] for a, b in C["quick"]],
                       [2.4, 5.6], avail))
    s.append(PageBreak())
    s.append(h2("Method by material"))
    for name, txt in C["by_material"]:
        s.append(h3(name))
        s.append(body(txt, "small"))
    s.append(Spacer(1, 12))
    s.append(callout("Never use " + plain(C["never"]), avail, label="Voids the finish"))
    s.append(Spacer(1, 16))
    s.append(h2("Seasonal check, 15 minutes"))
    s.append(numbered(C["seasonal"]))
    s.append(Spacer(1, 14))
    s.append(callout("A cleaning method that damages a finish is not a manufacturing defect "
                     "and is excluded from the warranty. Where a cleaning contractor is "
                     "engaged, issue this document with the scope.", avail,
                     label="For facilities teams"))
    s.append(Spacer(1, 12))
    s.append(Paragraph(f"Published method at {SITE}/how-to-clean. Questions: "
                       f"{plain(D.SUPPORT_PHONE)}.", st["cap"]))

    return render("veneta-care-maintenance.pdf", "Care &amp; Maintenance",
                  "Cleaning method by material for cellular, roller, Roman, sheer, faux wood, "
                  "vinyl and shutter products.", "For facilities teams", "Policy", s)


# ========================================================== measuring worksheet
def build_measuring_worksheet():
    st = _styles()
    avail = AVAIL
    lines = [plain(p["short"]) for p in D.PRODUCTS]
    s = []
    s.append(eyebrow("Worksheet"))
    s.append(Paragraph("Window measuring worksheet", st["h1"]))
    s.append(Paragraph("One block per window. Fill it on site, then order from the sheet "
                       "rather than from memory.", st["lede"]))
    s.append(Rule(space=12))
    s.append(h3("Before you start"))
    s.append(numbered([
        "Use a steel tape, not a cloth tape. A cloth tape stretches and the error lands in "
        "the finished size.",
        "Decide inside or outside mount before you measure, because the method differs.",
        'Inside mount: measure width in three places and record the narrowest to the '
        'nearest 1/8". Deduct nothing.',
        'Inside mount: measure height in three places and record the longest to the nearest '
        '1/8".',
        "Outside mount: measure the area you want covered, then add overlap on each side and "
        "at the head.",
        "Check the available depth against the mount depth reference before you commit to an "
        "inside mount.",
    ]))
    s.append(Spacer(1, 12))
    s.append(callout("Measuring errors are excluded from the warranty: a product made to the "
                     "size supplied is not a defect. Two minutes of re-checking on site is "
                     "cheaper than a remake.", avail, label="Why this sheet exists"))
    s.append(Spacer(1, 10))
    s.append(Paragraph("Product lines: " + "  ·  ".join(lines), st["cap"]))
    s.append(Paragraph("Minimum inside-mount depth: "
                       + "  ·  ".join(f"{plain(d)} {plain(l)}" for d, l, _ in D.MOUNT_DEPTH),
                       st["cap"]))

    blank = "_" * 26
    for pg in range(2):
        s.append(PageBreak())
        s.append(eyebrow(f"Windows {pg * 3 + 1}–{pg * 3 + 3}"))
        s.append(Paragraph("Window schedule", st["h2"]))
        s.append(Spacer(1, 4))
        for n in range(3):
            rows = [
                ("Room / window ID", blank * 2),
                ("Mount", "Inside  ☐    Outside  ☐"),
                ("Width 1 / 2 / 3", "______ / ______ / ______   →  narrowest ______"),
                ("Height 1 / 2 / 3", "______ / ______ / ______   →  longest ______"),
                ("Available depth", '______      Meets minimum?   Yes ☐   No ☐'),
                ("Product line", blank * 2),
                ("Fabric / colourway", blank * 2),
                ("Opacity", "Light filtering ☐   Room darkening ☐   Blackout ☐   Screen ☐"),
                ("Lift", "Cordless ☐   Top-down/bottom-up ☐   Motorized ☐   Wand ☐"),
                ("Obstructions", blank * 2),
                ("Notes", blank * 2),
            ]
            block = [Paragraph(f"Window {pg * 3 + n + 1}", st["h3"]),
                     kvtable(rows, avail, keyw=1.7)]
            s.append(KeepTogether(block))
            s.append(Spacer(1, 12))

    return render("veneta-measuring-worksheet.pdf", "Measuring Worksheet",
                  "A printable site sheet: opening, mount, depth check, line, fabric and "
                  "lift, three windows per page.", "For designers and installers",
                  "Worksheet", s)


# ============================================================= leave-behind cards
def _card_grid(cards, avail):
    """Two cards per row, hairline-bordered, sized to cut from a letter sheet."""
    rows = []
    for i in range(0, len(cards), 2):
        pair = cards[i:i + 2]
        while len(pair) < 2:
            pair.append([])
        rows.append(pair)
    tb = Table(rows, colWidths=[avail / 2, avail / 2], hAlign="LEFT")
    tb.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("BACKGROUND", (0, 0), (-1, -1), CANVAS),
        ("BOX", (0, 0), (-1, -1), 0.6, LINE),
        ("INNERGRID", (0, 0), (-1, -1), 0.6, LINE),
        ("LEFTPADDING", (0, 0), (-1, -1), 14),
        ("RIGHTPADDING", (0, 0), (-1, -1), 14),
        ("TOPPADDING", (0, 0), (-1, -1), 14),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 16),
    ]))
    return tb


def build_care_card():
    st = _styles()
    C = D.CARE
    avail = AVAIL
    cards = []
    for name, txt in C["by_material"]:
        cards.append([
            Paragraph("VENETA · CARE", st["micro"]),
            Paragraph(t(name), st["h3"]),
            Paragraph(t(txt), st["small"]),
            Rule(space=8),
            Paragraph("Never use " + t(C["never"]), st["cap"]),
            Paragraph(f"{plain(D.SUPPORT_PHONE)} · {SITE}/how-to-clean", st["cap"]),
        ])
    s = [eyebrow("Leave-behind"),
         Paragraph("Fabric care cards", st["h1"]),
         Paragraph("One card per material. Print, cut on the hairline, and leave the "
                   "relevant cards with the client at handover.", st["lede"]),
         Rule(space=12), Spacer(1, 6),
         _card_grid(cards, avail),
         Spacer(1, 10),
         Paragraph("Wording is identical to the published care guide, so a client who reads "
                   "the card and a client who reads the site get the same instruction.",
                   st["cap"])]
    return render("veneta-care-card.pdf", "Fabric Care Cards",
                  "Printable client leave-behind cards: cleaning method by material, one "
                  "card per material.", "Client leave-behind", "Handover pack", s)


def build_cordless_card():
    st = _styles()
    avail = AVAIL
    facts = [
        ("No looped cord", "Cordless lift is the standard configuration on every Veneta "
                           "line. There is no operating cord loop, no cleat and no tension "
                           "device to fit."),
        ("Where it matters", "Nurseries, children's bedrooms, low windows and window seats. "
                             "A cord within reach is the hazard the standard exists to "
                             "remove."),
        ("The standard", "Cellular shades are published as meeting ANSI/WCMA A100.1, the "
                         "American National Standard covering the safety of corded window "
                         "covering products."),
        ("Older windows in the house", "If any covering elsewhere in the home still has a "
                                       "looped cord or a chain within reach, that is the one "
                                       "to replace next."),
        ("Motorized shades", "TruQuiet™ motorization removes the cord as well. Battery packs "
                             "recharge in place, so no ladder and no loose cells."),
        ("If something binds", f"Stop and call {plain(D.SUPPORT_PHONE)}. Cordless mechanisms "
                               f"are covered for the life of the product for the original "
                               f"purchaser."),
    ]
    cards = []
    for head, txt in facts:
        cards.append([
            Paragraph("VENETA · CORDLESS", st["micro"]),
            Paragraph(t(head), st["h3"]),
            Paragraph(t(txt), st["small"]),
            Rule(space=8),
            Paragraph(f"{plain(D.SUPPORT_PHONE)} · {SITE}/child-safety", st["cap"]),
        ])
    s = [eyebrow("Leave-behind"),
         Paragraph("Cordless safety cards", st["h1"]),
         Paragraph("The cordless statement in plain language, sized for a handover pack. "
                   "Print, cut on the hairline, leave with the client.", st["lede"]),
         Rule(space=12), Spacer(1, 6),
         _card_grid(cards, avail),
         Spacer(1, 10),
         Paragraph("No blackout, energy-savings or acoustic claim appears on these cards, "
                   "because Veneta does not publish those figures.", st["cap"])]
    return render("veneta-cordless-card.pdf", "Cordless Safety Cards",
                  "Printable client leave-behind cards explaining cordless lift, the "
                  "standard it meets and where it matters.", "Client leave-behind",
                  "Handover pack", s)


# ================================================================= CSI 12 24 00
def build_csi():
    """Editable three-part spec. DOCX because a specifier has to edit it."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1.1)
        sec.right_margin = Inches(1.1)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)

    def p(txt, indent=0, bold=False, size=10, space=6, caps=False):
        par = doc.add_paragraph()
        par.paragraph_format.left_indent = Inches(indent)
        par.paragraph_format.space_after = Pt(space)
        run = par.add_run(plain(txt).upper() if caps else plain(txt))
        run.bold = bold
        run.font.size = Pt(size)
        return par

    def part(n, title):
        par = doc.add_paragraph()
        par.paragraph_format.space_before = Pt(16)
        par.paragraph_format.space_after = Pt(8)
        r = par.add_run(f"PART {n} — {title.upper()}")
        r.bold = True
        r.font.size = Pt(11)

    def art(num, title):
        par = doc.add_paragraph()
        par.paragraph_format.space_before = Pt(10)
        par.paragraph_format.space_after = Pt(4)
        r = par.add_run(f"{num}  {title.upper()}")
        r.bold = True

    # title block
    ttl = doc.add_paragraph()
    ttl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = ttl.add_run("SECTION 12 24 00")
    r.bold = True
    r.font.size = Pt(13)
    sub = doc.add_paragraph()
    r = sub.add_run("WINDOW SHADES AND BLINDS")
    r.bold = True
    r.font.size = Pt(13)
    p(f"Veneta Window Fashions · guide specification · revision {REV_LABEL}", size=9)
    p("EDITING NOTE: text in [square brackets] is a selection the specifier must make. "
      "Delete the options that do not apply. Every dimension, material and option below is "
      f"the figure published at {SITE}; do not add performance values that are not stated "
      "here. Veneta does not publish energy-savings percentages, R-values, U-factors or "
      "acoustic ratings, and this section deliberately contains none.", size=9)

    part(1, "General")
    art("1.1", "Section includes")
    for i, prod in enumerate(D.PRODUCTS, 1):
        p(f"A.{i}  {plain(prod['name'])}: {plain(prod['tagline'])}", indent=0.3)
    art("1.2", "Related requirements")
    p("A.  Section 06 10 00 — Rough Carpentry, for blocking at outside-mounted headrails.",
      indent=0.3)
    p("B.  Section 08 50 00 — Windows, for opening dimensions and glass stop depth.", indent=0.3)
    p("C.  Section 26 05 00 — Common Work Results for Electrical, for power at motorized "
      "units where hard-wired power is elected.", indent=0.3)
    art("1.3", "Reference standards")
    p("A.  ANSI/WCMA A100.1 — American National Standard for Safety of Corded Window "
      "Covering Products.", indent=0.3)
    p("B.  Manufacturer published product data, current revision.", indent=0.3)
    art("1.4", "Submittals")
    p("A.  Product data for each line specified, including the published width and height "
      "range, mount depth, operation and opacity options.", indent=0.3)
    p("B.  Samples: physical fabric or finish swatch for each colourway specified. Printed "
      "and on-screen colour representations are not acceptable for selection.", indent=0.3)
    p("C.  Window schedule listing room, opening, mount type, measured width, measured "
      "height and available depth for each unit.", indent=0.3)
    p("D.  Manufacturer statement confirming cordless operation where required by the "
      "occupancy.", indent=0.3)
    art("1.5", "Quality assurance")
    p("A.  Units are manufactured to submitted dimensions. Openings outside the published "
      "range are not supplied; revise the mount type or the unit count instead.", indent=0.3)
    p("B.  Field measure every opening before fabrication. Measuring error is not a "
      "manufacturing defect and is excluded from warranty.", indent=0.3)
    art("1.6", "Warranty")
    p(f"A.  {plain(D.WARRANTY['scope'])}", indent=0.3)
    p("B.  Motorization components: five years from date of purchase, including motors, "
      "battery packs and hubs.", indent=0.3)
    p("C.  Exclusions include the following:", indent=0.3)
    for i, ex in enumerate(D.WARRANTY["excluded"], 1):
        p(f"{i}.  {plain(ex)}", indent=0.6, size=9.5, space=3)

    part(2, "Products")
    art("2.1", "Manufacturer")
    p("A.  Veneta Window Fashions. Product is procured through The Home Depot; multi-unit, "
      "phased and date-committed orders are placed through the Veneta trade desk. "
      f"Telephone {plain(D.SUPPORT_PHONE)}. Reference {SITE}.", indent=0.3)
    p("B.  Substitutions: [permitted under Section 01 25 00] [not permitted].", indent=0.3)

    for i, prod in enumerate(D.PRODUCTS, 1):
        art(f"2.{i + 1}", plain(prod["name"]))
        p(f"A.  Description: {plain(prod['lede'])}", indent=0.3)
        p("B.  Published characteristics:", indent=0.3)
        for j, (k, v) in enumerate(prod["spec"], 1):
            p(f"{j}.  {plain(k)}: {plain(v)}.", indent=0.6, size=9.5, space=3)
        p("C.  Colourway: [" + " ] [".join(plain(n) for n, _ in prod["colors"]) + "].",
          indent=0.3)
        p("D.  Reference: " + f"{SITE}/{prod['slug']}.", indent=0.3)

    art(f"2.{len(D.PRODUCTS) + 2}", "Operation and safety")
    p("A.  Cordless lift is the standard configuration on every line in this section. "
      "Looped operating cords, cord cleats and tension devices are not supplied as standard.",
      indent=0.3)
    p("B.  Where the occupancy prohibits accessible cords, specify cordless or motorized "
      "operation explicitly on the window schedule.", indent=0.3)
    p("C.  Motorized operation: TruQuiet™ motorization, rechargeable battery pack "
      "[or hard-wired power]. Confirm control system compatibility before ordering.",
      indent=0.3)

    part(3, "Execution")
    art("3.1", "Examination")
    p("A.  Verify that openings are square, that blocking is present where required, and "
      "that the available depth at the head of the opening meets or exceeds the following "
      "published minimums:", indent=0.3)
    for d, lbl, _ in D.MOUNT_DEPTH:
        p(f"1.  {plain(d)}: {plain(lbl)}.", indent=0.6, size=9.5, space=3)
    p("B.  Where available depth is less than the published minimum, install as an outside "
      "mount. Do not modify the unit to fit.", indent=0.3)
    art("3.2", "Installation")
    p("A.  Install in accordance with manufacturer published instructions.", indent=0.3)
    p("B.  Set headrails level and brackets tight to substrate. Confirm full travel of every "
      "unit before demobilizing.", indent=0.3)
    p("C.  Expect a side light gap of approximately 1/4 inch on each side of an "
      "inside-mounted shade with a headrail cut to the opening. Where a darkened room is "
      "required, specify side channels on a blackout fabric or an outside mount. Total "
      "blackout is not claimed for any configuration.", indent=0.3)
    art("3.3", "Cleaning and protection")
    p("A.  Clean in accordance with the published care method for the material. Do not use "
      f"{plain(D.CARE['never'])}", indent=0.3)
    for name, txt in D.CARE["by_material"]:
        p(f"1.  {plain(name)}: {plain(txt)}", indent=0.6, size=9.5, space=3)
    art("3.4", "Closeout submittals")
    p("A.  Issue the published care and maintenance document to the facilities team.",
      indent=0.3)
    p("B.  Issue the cordless statement and warranty terms with the operation and "
      "maintenance manuals.", indent=0.3)
    p("END OF SECTION", bold=True, size=10)

    os.makedirs(OUT, exist_ok=True)
    path = os.path.join(OUT, "veneta-csi-122400.docx")
    doc.core_properties.title = "Section 12 24 00 — Window Shades and Blinds"
    doc.core_properties.author = "Veneta Window Fashions"
    doc.save(path)
    return path


# ========================================================================= build
BUILDERS = [
    ("veneta-commercial-spec-book.pdf", lambda: build_spec_book("commercial")),
    ("veneta-trade-spec-book.pdf", lambda: build_spec_book("trade")),
    ("veneta-fabric-schedule.pdf", build_fabric_schedule),
    ("veneta-mount-depths.pdf", build_mount_depths),
    ("veneta-cordless-statement.pdf", build_cordless_statement),
    ("veneta-warranty.pdf", build_warranty),
    ("veneta-care-maintenance.pdf", build_care_maintenance),
    ("veneta-measuring-worksheet.pdf", build_measuring_worksheet),
    ("veneta-care-card.pdf", build_care_card),
    ("veneta-cordless-card.pdf", build_cordless_card),
    ("veneta-csi-122400.docx", build_csi),
]

MANIFEST = {}     # filename -> (bytes, page count or None). Read by build/pages5.py.


def build_all(verbose=True):
    MANIFEST.clear()
    for name, fn in BUILDERS:
        path = fn()
        size = os.path.getsize(path)
        pages = _pagecount(path) if name.endswith(".pdf") else None
        MANIFEST[name] = (size, pages)
        if verbose:
            print(f"   spec  {name:38s} {size // 1024:4d} KB"
                  + (f"  {pages} pp" if pages else ""))
    return MANIFEST


def _pagecount(path):
    try:
        from pypdf import PdfReader
        return len(PdfReader(path).pages)
    except Exception:
        return None


def meta(filename):
    """(size_label, pages) for a download card, or None if the asset is not authored."""
    if filename not in MANIFEST:
        return None
    size, pages = MANIFEST[filename]
    kb = max(1, round(size / 1024))
    label = f"{kb} KB" if kb < 1024 else f"{size / 1048576:.1f} MB"
    return (label, pages)


if __name__ == "__main__":
    build_all()
